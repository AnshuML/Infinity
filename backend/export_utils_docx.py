from io import BytesIO
from docx import Document

def scope_to_docx(scope):
    doc = Document()
    doc.add_heading('PROPOSED SCOPE DOCUMENT', level=1)
    doc.add_heading('1. Project Overview', level=2)
    doc.add_paragraph(f"Title: {scope.project_title}")
    doc.add_heading('2. Requirements Summary', level=2)
    doc.add_heading('3. Scope of Work', level=2)
    doc.add_paragraph('In-Scope:')
    for s in scope.scope_in:
        doc.add_paragraph(f"- {s}")
    doc.add_paragraph('Out-of-Scope:')
    for s in scope.scope_out:
        doc.add_paragraph(f"- {s}")
    doc.add_heading('4. Navigation & Sitemap', level=2)
    for n in scope.navigation:
        doc.add_paragraph(f"- {n}")
    doc.add_heading('6. Feedback & Approval', level=2)
    for gap in scope.gap_analysis:
        doc.add_paragraph(f"> {gap}")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def framework_to_docx(framework):
    doc = Document()
    doc.add_heading('CONTENT FRAMEWORK', level=1)
    
    doc.add_heading('1. Header Navigation', level=2)
    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Main Nav'
    hdr_cells[1].text = 'Dropdown'
    hdr_cells[2].text = 'Destination'
    hdr_cells[3].text = 'Type'
    hdr_cells[4].text = 'Description'
    hdr_cells[5].text = 'Content Type'
    
    for item in framework.header_nav:
        row_cells = table.add_row().cells
        row_cells[0].text = item.main_nav
        row_cells[1].text = item.dropdown
        row_cells[2].text = item.final_destination
        row_cells[3].text = item.page_type
        row_cells[4].text = item.page_description
        row_cells[5].text = item.content_type

    doc.add_heading('2. Footer Navigation', level=2)
    table2 = doc.add_table(rows=1, cols=5)
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Menu Title'
    hdr_cells2[1].text = 'Nested Items'
    hdr_cells2[2].text = 'Type'
    hdr_cells2[3].text = 'Description'
    hdr_cells2[4].text = 'Content Type'
    
    for item in framework.footer_nav:
        row_cells = table2.add_row().cells
        row_cells[0].text = item.menu_title
        row_cells[1].text = item.nested_items
        row_cells[2].text = item.page_type
        row_cells[3].text = item.page_description
        row_cells[4].text = item.content_type

    doc.add_heading('3. Website Assets', level=2)
    table3 = doc.add_table(rows=1, cols=3)
    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = 'Asset Required'
    hdr_cells3[1].text = 'Description'
    hdr_cells3[2].text = 'Content Type'
    
    for item in framework.website_assets:
        row_cells = table3.add_row().cells
        row_cells[0].text = item.asset_required
        row_cells[1].text = item.description
        row_cells[2].text = item.content_type

    doc.add_heading('CTA Strategy', level=2)
    doc.add_paragraph(framework.cta_strategy)
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

